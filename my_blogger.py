import os
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
import google.generativeai as genai
from typing import List, Dict
from pydantic import BaseModel, Field
import json
import uuid
from pydantic import ValidationError
from docx import Document
from fastapi import FastAPI, Query
import uvicorn

load_dotenv()

app = FastAPI(title="Trending Blog Generator API")

# 🔑 Google Custom Search API keys
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GOOGLE_CX = os.getenv("GOOGLE_CX")
genai.configure(api_key=GOOGLE_API_KEY)

class BlogSection(BaseModel):
    heading: str = Field(..., description="Heading of the blog section")
    content: str = Field(..., description="Content/paragraph text for this section")

class BlogOutput(BaseModel):
    title: str = Field(..., description="Short, engaging blog title (<=12 words)")
    subtitle: str = Field(..., description="Catchy subtitle (<=20 words)")
    body: List[BlogSection] = Field(default_factory=list, description="List of blog sections with headings and content")
    conclusion: str = Field("", description="Closing summary paragraph (~10% of length)")
    tags: List[str] = Field(default_factory=list, description="Top 5 keyword tags for the blog")
    seo_meta: str = Field("", description="SEO meta description (<=160 characters)")

class TopicDetails(BaseModel):
    title: str = Field(..., description="Chosen blog title")
    url: str = Field(..., description="Corresponding blog URL")

search_store: Dict[str, List[Dict[str, str]]] = {}


def google_search(query, num_results=10):
    """Search Google with Custom Search API"""
    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "q": query,
        "key": GOOGLE_API_KEY,
        "cx": GOOGLE_CX,
        "num": num_results,
    }
    res = requests.get(url, params=params)
    res.raise_for_status()
    return res.json().get("items", [])

def choose_topic_auto(topics: List[Dict[str, str]]) -> TopicDetails:
    prompt = ChatPromptTemplate.from_template(
        """
            Here are some trending blog titles and links:
            {topics}

            Please select the single most trending and relevant blog.
            Return the result strictly in JSON format like this:
            {{
                "title": "chosen blog title",
                "url": "corresponding url"
            }}

            Note: Don't choose "https://chatgpt.com", "https://chat.com", alternatives kind of domain as it requires signin.
        """
    )
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", api_key=GOOGLE_API_KEY).with_structured_output(TopicDetails)
    try:
        chain = prompt | model
        chosen_topic = chain.invoke({"topics": topics})
        return chosen_topic
    except Exception as e:
        print(f"❌ Auto choose failed: {e}")
        fallback = topics[0]
        return TopicDetails(title=fallback["title"], url=fallback["url"])

def fetch_blog_text(url: str) -> str:
    try:
        html = requests.get(url, timeout=10).text
        soup = BeautifulSoup(html, "html.parser")
        # crude extraction: remove scripts/styles
        for tag in soup(["script", "style"]):
            tag.extract()
        text = " ".join(soup.stripped_strings)
        return text[:4000]  # limit length for model
    except Exception as e:
        return f"Error fetching {url}: {e}"

def summarize_blog(url: str) -> str:
    text = fetch_blog_text(url)
    prompt=f"Summarize this blog in 5 sentences:\n\n{text}"
    response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
    return response.text.strip()


def generate_blog(topic: str, summary: str = "", tone: str = "professional", word_target: int = 700) -> BlogOutput:
    print("::::::::::", topic, summary, word_target, tone)
    prompt = ChatPromptTemplate.from_template(
        """
        You are an expert content writer. Write clear, engaging, and original long-form blogs.

        Generate a long-form blog post ~{word_target} words about this topic: '{topic}'.
        context: {summary}
        Tone: {tone}.

        Requirements:
        - Provide output in JSON with keys: title, subtitle, body, seo_meta
        - Title: short, engaging (<=12 words)
        - Subtitle: catchy (<=20 words)
        - Body: split into multiple sections with headings
        - Tags: top five key for this blog
        - SEO Meta: concise description (<=160 chars)

        Format:
        {{
            "title": "short, engaging (<=12 words)",
            "subtitle": "catchy (<=20 words)",
            "body": [
                {{"heading": "Section Heading", "content": "Paragraph text"}},
                ...
            ],
            "conclusion": "typically makes up around 10% of the total document's word count, or about 5-7% for theses and dissertations.",
            "tags": ["keyword1", "keyword2"],
            "seo_meta": "meta description (<=160 chars)"
        }}

        Return valid JSON ONLY.
        """
        )
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", api_key=GOOGLE_API_KEY).with_structured_output(BlogOutput)
    try:
        chain = prompt | model
        print(f"📝 Generating blog for topic: {topic}")
        blog = chain.invoke({"topic": topic, "summary": summary, "tone": tone, "word_target": word_target})
        print("::::::::::::::::::::::::::", blog)
        return blog
    except ValidationError as e:
        print("❌ Validation failed:", e)
        raise
    except Exception:
        return BlogOutput(
            title=topic,
            subtitle="",
            body=[{"heading": "Generated Content", "content": "Blog generation failed, please retry."}],
            tags=[],
            seo_meta=""
        )

def export_blog_to_docx(blog, filename):
    doc = Document()

    # Title + Subtitle
    doc.add_heading(blog.title, 0)
    if blog.subtitle:
        doc.add_paragraph(blog.subtitle, style="Subtitle")

    doc.add_paragraph("")  # spacer

    # Body sections
    for section in blog.body:
        if section.heading:
            doc.add_heading(section.heading, 2)
        if section.content:
            doc.add_paragraph(section.content)

        doc.add_paragraph("")  # spacer between sections

    # Conclusion (if exists)
    if hasattr(blog, "conclusion") and blog.conclusion:
        doc.add_heading("Conclusion", 2)
        doc.add_paragraph(blog.conclusion)
        doc.add_paragraph("")

    # SEO Meta
    if blog.seo_meta:
        p = doc.add_paragraph(blog.seo_meta)
        run = p.runs[0]
        run.italic = True
        doc.add_paragraph("")  # spacer

    # Tags
    if blog.tags:
        doc.add_heading("Tags", 2)
        for tag in blog.tags:
            doc.add_paragraph(tag, style="List Bullet")

    safe_filename = blog.title.replace(" ", "_").replace("/", "_").replace(":", "_")
    filename = safe_filename + ".docx"
    doc.save(filename)
    print(f"✅ Blog exported to {filename}")


# @app.get("/search")
# def search_blogs(query: str = Query(...)):
#     """Step 1: Google search blogs"""
#     results = google_search(query)
#     titles = [{"title": r.get("title"), "url": r.get("link")} for r in results]
#     search_id = str(uuid.uuid4())
#     search_store[search_id] = titles

#     return {"search_id": search_id, "results": titles}

# @app.get("/choose/auto")
# def choose_auto(search_id: str):
#     """Automatically pick the best topic using Gemini"""
#     topics = search_store.get(search_id)
#     if not topics:
#         return {"error": "Invalid or expired search_id"}
    
#     chosen_title = choose_topic_auto(topics)

#     # find corresponding URL
#     chosen_item = next((t for t in topics if t["title"] == chosen_title), None)
#     if not chosen_item:
#         return {"error": "Gemini’s chosen title not found in results"}

#     return {"search_id": search_id, "chosen_topic": chosen_item}


# @app.get("/choose/human")
# def choose_human(search_id: str, index: int = Query(..., description="Index of title from /search response")):
#     """Manually pick a topic by index"""
#     topics = search_store.get(search_id)
#     if not topics:
#         return {"error": "Invalid or expired search_id"}
    
#     if index < 1 or index > len(topics):
#         return {"error": f"Invalid index. Choose between 1 and {len(topics)}"}
    
#     chosen_item = topics[index - 1]
#     return {"search_id": search_id, "chosen_topic": chosen_item}



# @app.get("/summarize")
# def summarize_url(url: str = Query(...)):
#     """Step 3: Summarize one blog from URL"""
#     return {"url": url, "summary": summarize_blog(url)}


@app.get("/generate-blog")
def generate_blog_endpoint(query: str = Query(..., description="Search query for trending blogs"), tone: str = Query("professional", description="Tone of the blog (e.g., professional, casual, persuasive)"),
    word_target: int = Query(700, description="Target word count for the blog"),
):
    results = google_search(query)
    topics = [{"title": r.get("title"), "url": r.get("link")} for r in results]
    print("------------------topics===", topics)
    if not topics:
        return {"error": "No search results found"}

    chosen_item = choose_topic_auto(topics)
    if not chosen_item or not chosen_item.title:
        return {"error": "Gemini did not return a valid topic"}

    summary = summarize_blog(chosen_item.url)

    blog = generate_blog(chosen_item.title, summary, tone, word_target)

    safe_filename = blog.title.replace(" ", "_").replace("/", "_")
    filepath = f"{safe_filename}.docx"
    export_blog_to_docx(blog, safe_filename)

    return {
        "chosen_topic": chosen_item,
        "summary": summary,
        "docx_file": filepath
    }

if __name__ == '__main__':
    uvicorn.run(app, host="0.0.0.0", port=8000)