import streamlit as st
import requests
import json
import os
from datetime import datetime
import base64

# Streamlit app configuration
st.set_page_config(
    page_title="Trending Blog Generator",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# if "disabled" not in st.session_state:
#     st.session_state.disabled = False

# Custom CSS styling
st.markdown("""
    <style>
    .main-header {
        font-size: 3rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .blog-card {
        background-color: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 5px solid #1f77b4;
        margin-bottom: 1rem;
    }
    .section-header {
        color: #1f77b4;
        border-bottom: 2px solid #1f77b4;
        padding-bottom: 0.5rem;
        margin-top: 1.5rem;
    }
    .stButton button {
        background-color: #1f77b4;
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 5px;
        width: 100%;
    }
    .stButton button:hover {
        background-color: #0d6efd;
    }
    </style>
""", unsafe_allow_html=True)

class BlogGeneratorUI:
    def __init__(self):
        self.api_url = "http://localhost:8000"  # FastAPI server URL
        
    def create_download_link(self, file_path):
        """Create a download link for the generated DOCX file"""
        try:
            with open(file_path, "rb") as file:
                data = file.read()
            b64 = base64.b64encode(data).decode()
            filename = os.path.basename(file_path)
            href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">📥 Download DOCX</a>'
            return href
        except FileNotFoundError:
            return "File not found yet. Please wait for generation to complete."
    
    def display_blog_content(self, blog_data):
        """Display and allow editing of generated blog content"""
        if "error" in blog_data:
            st.error(f"Error: {blog_data['error']}")
            return

        st.markdown(f"### 📖 Generated Blog: {blog_data['chosen_topic']['title']}")

        # Editable blog content
        if 'blog_content' in blog_data:
            blog = blog_data['blog_content']

            st.subheader("✏️ Edit Your Blog")
            edited_title = st.text_input("Blog Title", blog['title'])
            edited_subtitle = st.text_area("Subtitle", blog['subtitle'])

            edited_sections = []
            for i, section in enumerate(blog['body']):
                st.markdown(f"#### Section {i+1}")
                heading = st.text_input(f"Heading {i+1}", section['heading'])
                content = st.text_area(f"Content {i+1}", section['content'], height=150)
                edited_sections.append({"heading": heading, "content": content})

            edited_conclusion = st.text_area("Conclusion", blog.get("conclusion", ""))
            edited_tags = st.text_input("Tags (comma separated)", ", ".join(blog.get("tags", [])))
            edited_seo = st.text_area("SEO Meta Description", blog.get("seo_meta", ""))

            if st.button("📥 Export as DOCX"):
                from docx import Document
                doc = Document()

                doc.add_heading(edited_title, 0)
                if edited_subtitle:
                    doc.add_paragraph(edited_subtitle, style="Intense Quote")

                for section in edited_sections:
                    doc.add_heading(section["heading"], level=1)
                    doc.add_paragraph(section["content"])

                if edited_conclusion:
                    doc.add_heading("Conclusion", level=1)
                    doc.add_paragraph(edited_conclusion)

                if edited_tags:
                    doc.add_heading("Tags", level=2)
                    doc.add_paragraph(edited_tags)

                if edited_seo:
                    p = doc.add_paragraph(edited_seo)
                    if p.runs:
                        p.runs[0].italic = True

                file_path = f"blog_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                doc.save(file_path)
                st.success("✅ Blog exported successfully!")
                st.markdown(self.create_download_link(file_path), unsafe_allow_html=True)

    
    def generate_blog(self, query, tone, word_target):
        """Call the FastAPI endpoint to generate blog"""
        try:
            response = requests.get(
                f"{self.api_url}/generate-blog",
                params={
                    "query": query,
                    "tone": tone,
                    "word_target": word_target
                },
                timeout=120  # 2 minutes timeout for generation
            )
            
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"API Error: {response.status_code} - {response.text}"}
                
        except requests.exceptions.RequestException as e:
            return {"error": f"Connection error: {str(e)}"}
        except json.JSONDecodeError:
            return {"error": "Invalid response from server"}
    
    def render_sidebar(self):
        """Render the sidebar with input controls"""
        with st.sidebar:
            st.markdown("## ⚙️ Blog Generation Settings")
            
            # Search query input
            query = st.text_input(
                "🔍 Search Query",
                value="latest technology trends 2025",
                help="Enter a search term to find trending topics"
                # disabled=st.session_state.disabled
            )
            
            # Tone selection
            tone = st.selectbox(
                "🎭 Writing Tone",
                options=["professional", "casual", "persuasive", "informative", "enthusiastic"],
                index=0,
                help="Select the tone for your blog post"
                # disabled=st.session_state.is_generating
            )
            
            # Word target input
            word_target = st.slider(
                "📏 Target Word Count",
                min_value=300,
                max_value=2000,
                value=700,
                step=100,
                help="Approximate word count for the generated blog"
                # disabled=st.session_state.is_generating
            )
            
            # Advanced options
            with st.expander("Advanced Options"):
                st.checkbox("Include source analysis", value=True)
                st.checkbox("Generate multiple variants", value=False)
            
            # Generate button
            generate_btn = st.button(
                "🚀 Generate Blog",
                type="primary",
                use_container_width=True
                # disabled=st.session_state.disabled
            )
            
            st.markdown("---")
            st.markdown("### 📊 Recent Activity")
            # Placeholder for recent activity log
            st.info("No recent generations yet")
            
            st.markdown("---")
            st.markdown("### ℹ️ About")
            st.caption("""
                This tool uses AI to generate high-quality blog posts 
                based on trending topics from Google search results.
                Powered by Gemini AI and FastAPI.
            """)
            
            return query, tone, word_target, generate_btn
    
    def render_main_content(self):
        """Render the main content area"""
        st.markdown('<h1 class="main-header">📝 AI Blog Generator</h1>', unsafe_allow_html=True)
        st.markdown("Generate engaging blog posts from trending topics using AI")

        # Sidebar inputs
        query, tone, word_target, generate_btn = self.render_sidebar()

        # Tabs
        tab1, tab2, tab3 = st.tabs(["🎯 Generate", "📚 Examples", "ℹ️ Guide"])

        with tab1:
            col1, col2 = st.columns([2, 1])

            with col2:
                st.markdown("### ⚡ Quick Actions")
                if st.button("🔄 Test Connection", use_container_width=True):
                    try:
                        response = requests.get(f"{self.api_url}/", timeout=5)
                        if response.status_code == 200:
                            st.success("✅ API Connected successfully!")
                        else:
                            st.error("❌ API connection failed")
                    except:
                        st.error("❌ Cannot connect to API server")

                st.markdown("### 🎯 Current Settings")
                st.info(f"**Search:** {query}")
                st.info(f"**Tone:** {tone}")
                st.info(f"**Word Target:** {word_target} words")

            with col1:
                if generate_btn:
                    if not query.strip():
                        st.error("Please enter a search query")
                        return

                    # st.session_state.disabled = True
                    with st.spinner("🔍 Searching for trending topics..."):
                        # (Here you’d call API for trending topics if needed)
                        pass
                    

                    with st.spinner("Generating your blog post... This may take a few minutes"):
                        result = self.generate_blog(query, tone, word_target)

                        # Display editable blog
                        self.display_blog_content(result)
                    # st.session_state.disabled = False
        with tab2:
            st.markdown("### 📚 Example Use Cases")
            
            examples = [
                {
                    "query": "artificial intelligence healthcare applications",
                    "tone": "professional",
                    "description": "Technical overview of AI in medical field"
                },
                {
                    "query": "sustainable fashion trends",
                    "tone": "casual",
                    "description": "Trendy article about eco-friendly fashion"
                },
                {
                    "query": "digital marketing strategies 2025",
                    "tone": "persuasive",
                    "description": "Convincing piece about modern marketing"
                }
            ]
            
            for example in examples:
                with st.expander(f"🔍 {example['query']}"):
                    st.write(f"**Tone:** {example['tone']}")
                    st.write(example['description'])
                    if st.button("Use this example", key=example['query']):
                        st.session_state.query = example['query']
                        st.session_state.tone = example['tone']
                        st.rerun()
        
        with tab3:
            st.markdown("### 📖 User Guide")
            st.markdown("""
            #### How to use this tool:
            
            1. **Enter a search query** - What topic do you want to write about?
            2. **Choose a tone** - Professional, casual, persuasive, etc.
            3. **Set word target** - How long should the blog be?
            4. **Click Generate** - Wait for the AI to work its magic!
            
            #### Features:
            - ✅ Automatic topic research from Google trends
            - ✅ AI-powered content generation
            - ✅ Multiple writing tones
            - ✅ Word document export
            - ✅ SEO optimization
            
            #### Tips:
            - Be specific with your search queries for better results
            - Use professional tone for technical topics
            - 700-1000 words is ideal for most blog posts
            """)
    
    def run(self):
        """Run the Streamlit application"""
        try:
            self.render_main_content()
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")
            st.info("Please make sure the FastAPI server is running on http://localhost:8000")

# Run the application
if __name__ == "__main__":
    # Check if API server is running
    try:
        response = requests.get("http://localhost:8000/", timeout=2)
        if response.status_code != 200:
            st.warning("⚠️ FastAPI server might not be running. Please start it with: `uvicorn main:app --reload`")
    except:
        st.warning("⚠️ Cannot connect to FastAPI server. Please make sure it's running on http://localhost:8000")
    
    # Initialize and run the UI
    app = BlogGeneratorUI()
    app.run()