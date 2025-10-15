# Human in the loop
# Multiple blogs blending

# Add ai content after adding new content box [human will add the heading rest if the content use button(generate with ai)], heading refresh with 2-3 options
# AI Generated content checker

# Image multiple with any white/black background check insta for reference
# Image Generation with category/style

# Image with quote
# Image with Film quote with relatable film image(quote on the generate/webscrap: for image)

# Custom image bg, suggestions new columns

import os
import re
import requests
import sqlite3
from datetime import datetime
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
import google.generativeai as genai
from typing import List, Dict
from pydantic import BaseModel, Field
import json
import uuid
from pydantic import ValidationError
from docx import Document
from fastapi import FastAPI, Query, HTTPException
import uvicorn
from fastapi.middleware.cors import CORSMiddleware
import asyncio
from pathlib import Path
from instaImages import generate_instagram_carousel, extract_slides, generate_slide_image

load_dotenv()

app = FastAPI(title="Trending Blog Generator API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For development only, restrict in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

blogs_folder = Path("blogs")
blogs_folder.mkdir(exist_ok=True)

# Create Instagram images folder
instagram_folder = Path("instagram_images")
instagram_folder.mkdir(exist_ok=True)

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GOOGLE_CX = os.getenv("GOOGLE_CX")
genai.configure(api_key=GOOGLE_API_KEY)

# Database setup
DB_PATH = "blog_generator.db"

def init_db():
    """Initialize SQLite database with required tables"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Create blogs table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS blogs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            search_id TEXT,
            title TEXT NOT NULL,
            subtitle TEXT,
            content_json TEXT NOT NULL,
            summary TEXT,
            linkedin_version TEXT,
            linkedin_content TEXT,
            docx_filename TEXT,
            tone TEXT DEFAULT 'professional',
            word_count INTEGER DEFAULT 700,
            tags TEXT,
            seo_meta TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create search_sessions table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS search_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            search_id TEXT UNIQUE NOT NULL,
            query TEXT NOT NULL,
            topics_json TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create blended_blogs table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS blended_blogs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            search_id TEXT,
            title TEXT NOT NULL,
            subtitle TEXT,
            content_json TEXT NOT NULL,
            blended_summary TEXT,
            linkedin_version TEXT,
            linkedin_content TEXT,
            docx_filename TEXT,
            topic_indices TEXT,
            individual_summaries TEXT,
            tone TEXT DEFAULT 'professional',
            word_count INTEGER DEFAULT 700,
            tags TEXT,
            seo_meta TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Create instagram_posts table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS instagram_posts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            blog_id INTEGER,
            blog_type TEXT DEFAULT 'regular', -- 'regular' or 'blended'
            carousel_text TEXT,
            images_json TEXT, -- JSON array of image filenames
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (blog_id) REFERENCES blogs(id) ON DELETE CASCADE
        )
    ''')
    
    # Create linkedin_posts table for better LinkedIn content management
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS linkedin_posts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            blog_id INTEGER,
            blog_type TEXT DEFAULT 'regular', -- 'regular' or 'blended'
            post_content TEXT,
            hashtags TEXT,
            character_count INTEGER,
            engagement_score REAL, -- Can be calculated based on content quality
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (blog_id) REFERENCES blogs(id) ON DELETE CASCADE
        )
    ''')
    
    conn.commit()
    conn.close()

def save_search_session(search_id: str, query: str, topics: List[Dict[str, str]]):
    """Save search session to database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT OR REPLACE INTO search_sessions (search_id, query, topics_json)
        VALUES (?, ?, ?)
    ''', (search_id, query, json.dumps(topics)))
    
    conn.commit()
    conn.close()

def get_search_session(search_id: str) -> List[Dict[str, str]]:
    """Retrieve search session from database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('SELECT topics_json FROM search_sessions WHERE search_id = ?', (search_id,))
    result = cursor.fetchone()
    conn.close()
    
    if result:
        return json.loads(result[0])
    return None

def save_blog_outcome(
    search_id: str,
    title: str,
    subtitle: str,
    content_json: dict,
    summary: str,
    linkedin_version: str,
    linkedin_content: str,
    docx_filename: str,
    tone: str,
    word_count: int,
    tags: List[str],
    seo_meta: str
):
    """Save blog outcome to database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO blogs (
            search_id, title, subtitle, content_json, summary, linkedin_version,
            linkedin_content, docx_filename, tone, word_count, tags, seo_meta
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        search_id,
        title,
        subtitle,
        json.dumps(content_json),
        summary,
        linkedin_version,
        linkedin_content,
        docx_filename,
        tone,
        word_count,
        json.dumps(tags) if tags else None,
        seo_meta
    ))
    
    blog_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return blog_id

def save_blended_blog_outcome(
    search_id: str,
    title: str,
    subtitle: str,
    content_json: dict,
    blended_summary: str,
    linkedin_version: str,
    linkedin_content: str,
    docx_filename: str,
    topic_indices: List[int],
    individual_summaries: List[str],
    tone: str,
    word_count: int,
    tags: List[str],
    seo_meta: str
):
    """Save blended blog outcome to database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO blended_blogs (
            search_id, title, subtitle, content_json, blended_summary, linkedin_version,
            linkedin_content, docx_filename, topic_indices, individual_summaries, tone, word_count, tags, seo_meta
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        search_id,
        title,
        subtitle,
        json.dumps(content_json),
        blended_summary,
        linkedin_version,
        linkedin_content,
        docx_filename,
        json.dumps(topic_indices),
        json.dumps(individual_summaries),
        tone,
        word_count,
        json.dumps(tags) if tags else None,
        seo_meta
    ))
    
    blog_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return blog_id

def save_instagram_post(blog_id: int, blog_type: str, carousel_text: str, images: List[str]):
    """Save Instagram carousel data to database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO instagram_posts (blog_id, blog_type, carousel_text, images_json)
        VALUES (?, ?, ?, ?)
    ''', (blog_id, blog_type, carousel_text, json.dumps(images)))
    
    conn.commit()
    conn.close()

def save_linkedin_post(blog_id: int, blog_type: str, post_content: str, hashtags: str = None):
    """Save LinkedIn post data to database"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Extract hashtags from content if not provided
    if not hashtags:
        hashtag_matches = re.findall(r'#\w+', post_content)
        hashtags = ' '.join(hashtag_matches) if hashtag_matches else ''
    
    character_count = len(post_content)
    
    # Simple engagement score calculation (can be enhanced)
    engagement_score = min(100, (character_count / 10) + (len(hashtag_matches) * 5))
    
    cursor.execute('''
        INSERT INTO linkedin_posts (blog_id, blog_type, post_content, hashtags, character_count, engagement_score)
        VALUES (?, ?, ?, ?, ?, ?)
    ''', (blog_id, blog_type, post_content, hashtags, character_count, engagement_score))
    
    conn.commit()
    conn.close()

def get_all_blogs(limit: int = 50, offset: int = 0):
    """Retrieve all blogs with pagination"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT id, search_id, title, subtitle, docx_filename, tone, word_count, 
               linkedin_version, created_at
        FROM blogs 
        ORDER BY created_at DESC 
        LIMIT ? OFFSET ?
    ''', (limit, offset))
    
    blogs = []
    for row in cursor.fetchall():
        blogs.append({
            'id': row[0],
            'search_id': row[1],
            'title': row[2],
            'subtitle': row[3],
            'docx_filename': row[4],
            'tone': row[5],
            'word_count': row[6],
            'linkedin_version': row[7],
            'created_at': row[8]
        })
    
    conn.close()
    return blogs

def get_blog_by_id(blog_id: int):
    """Retrieve complete blog data by ID"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT * FROM blogs WHERE id = ?
    ''', (blog_id,))
    
    row = cursor.fetchone()
    if not row:
        conn.close()
        return None
    
    # Get Instagram posts for this blog
    cursor.execute('''
        SELECT carousel_text, images_json, created_at 
        FROM instagram_posts 
        WHERE blog_id = ? AND blog_type = 'regular'
        ORDER BY created_at DESC
    ''', (blog_id,))
    
    instagram_posts = []
    for insta_row in cursor.fetchall():
        instagram_posts.append({
            'carousel_text': insta_row[0],
            'images': json.loads(insta_row[1]) if insta_row[1] else [],
            'created_at': insta_row[2]
        })
    
    # Get LinkedIn posts for this blog
    cursor.execute('''
        SELECT post_content, hashtags, character_count, engagement_score, created_at 
        FROM linkedin_posts 
        WHERE blog_id = ? AND blog_type = 'regular'
        ORDER BY created_at DESC
    ''', (blog_id,))
    
    linkedin_posts = []
    for linkedin_row in cursor.fetchall():
        linkedin_posts.append({
            'post_content': linkedin_row[0],
            'hashtags': linkedin_row[1],
            'character_count': linkedin_row[2],
            'engagement_score': linkedin_row[3],
            'created_at': linkedin_row[4]
        })
    
    blog = {
        'id': row[0],
        'search_id': row[1],
        'title': row[2],
        'subtitle': row[3],
        'content_json': json.loads(row[4]),
        'summary': row[5],
        'linkedin_version': row[6],
        'linkedin_content': row[7],
        'docx_filename': row[8],
        'tone': row[9],
        'word_count': row[10],
        'tags': json.loads(row[11]) if row[11] else [],
        'seo_meta': row[12],
        'created_at': row[13],
        'updated_at': row[14],
        'instagram_posts': instagram_posts,
        'linkedin_posts': linkedin_posts
    }
    
    conn.close()
    return blog

def get_blended_blogs(limit: int = 50, offset: int = 0):
    """Retrieve all blended blogs with pagination"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT id, search_id, title, subtitle, docx_filename, tone, word_count, 
               linkedin_version, created_at
        FROM blended_blogs 
        ORDER BY created_at DESC 
        LIMIT ? OFFSET ?
    ''', (limit, offset))
    
    blogs = []
    for row in cursor.fetchall():
        blogs.append({
            'id': row[0],
            'search_id': row[1],
            'title': row[2],
            'subtitle': row[3],
            'docx_filename': row[4],
            'tone': row[5],
            'word_count': row[6],
            'linkedin_version': row[7],
            'created_at': row[8]
        })
    
    conn.close()
    return blogs

def get_blended_blog_by_id(blog_id: int):
    """Retrieve complete blended blog data by ID"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT * FROM blended_blogs WHERE id = ?
    ''', (blog_id,))
    
    row = cursor.fetchone()
    if not row:
        conn.close()
        return None
    
    # Get Instagram posts for this blended blog
    cursor.execute('''
        SELECT carousel_text, images_json, created_at 
        FROM instagram_posts 
        WHERE blog_id = ? AND blog_type = 'blended'
        ORDER BY created_at DESC
    ''', (blog_id,))
    
    instagram_posts = []
    for insta_row in cursor.fetchall():
        instagram_posts.append({
            'carousel_text': insta_row[0],
            'images': json.loads(insta_row[1]) if insta_row[1] else [],
            'created_at': insta_row[2]
        })
    
    # Get LinkedIn posts for this blended blog
    cursor.execute('''
        SELECT post_content, hashtags, character_count, engagement_score, created_at 
        FROM linkedin_posts 
        WHERE blog_id = ? AND blog_type = 'blended'
        ORDER BY created_at DESC
    ''', (blog_id,))
    
    linkedin_posts = []
    for linkedin_row in cursor.fetchall():
        linkedin_posts.append({
            'post_content': linkedin_row[0],
            'hashtags': linkedin_row[1],
            'character_count': linkedin_row[2],
            'engagement_score': linkedin_row[3],
            'created_at': linkedin_row[4]
        })
    
    blog = {
        'id': row[0],
        'search_id': row[1],
        'title': row[2],
        'subtitle': row[3],
        'content_json': json.loads(row[4]),
        'blended_summary': row[5],
        'linkedin_version': row[6],
        'linkedin_content': row[7],
        'docx_filename': row[8],
        'topic_indices': json.loads(row[9]) if row[9] else [],
        'individual_summaries': json.loads(row[10]) if row[10] else [],
        'tone': row[11],
        'word_count': row[12],
        'tags': json.loads(row[13]) if row[13] else [],
        'seo_meta': row[14],
        'created_at': row[15],
        'updated_at': row[16],
        'instagram_posts': instagram_posts,
        'linkedin_posts': linkedin_posts
    }
    
    conn.close()
    return blog

# Initialize database on startup
init_db()

class BlogSection(BaseModel):
    heading: str = Field(..., description="Heading of the blog section")
    content: str = Field(..., description="Content/paragraph text for this section")

class BlogOutput(BaseModel):
    title: str = Field(..., description="Short, engaging blog title (<=12 words)")
    subtitle: str = Field(..., description="Catchy subtitle (<=20 words)")
    body: List[BlogSection] = Field(default_factory=list, description="List of blog sections with headings and content")
    conclusion: str = Field("", description="Closing summary paragraph (~10% of length)")
    tags: List[str] = Field(default_factory=list, description="Top 5 keyword tags for the blog")
    seo_meta: str = Field("", description="SEO meta description (<=160 characters)")

class TopicDetails(BaseModel):
    title: str = Field(..., description="Chosen blog title")
    url: str = Field(..., description="Corresponding blog URL")

class MultiBlogRequest(BaseModel):
    topic_indices: List[int] = Field(..., description="List of topic indices to blend")
    search_id: str = Field(..., description="Search session ID")
    tone: str = Field("professional", description="Tone of the blog")
    word_target: int = Field(700, description="Target word count")

class InstagramRequest(BaseModel):
    blog_content: str
    blog_id: int = Field(None, description="ID of the blog to associate with")
    blog_type: str = Field("regular", description="Type of the blog")
    customization_prompt: str | None = "Use clean white background with minimal design"

search_store: Dict[str, List[Dict[str, str]]] = {}


def google_search(query, num_results=10):
    """Search Google with Custom Search API"""
    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "q": query,
        "key": GOOGLE_API_KEY,
        "cx": GOOGLE_CX,
        "num": num_results,
    }
    res = requests.get(url, params=params)
    res.raise_for_status()
    return res.json().get("items", [])

def choose_topic_auto(topics: List[Dict[str, str]]) -> TopicDetails:
    prompt = ChatPromptTemplate.from_template(
        """
            Here are some trending blog titles and links:
            {topics}

            Please select the single most trending and relevant blog.
            Return the result strictly in JSON format like this:
            {{
                "title": "chosen blog title",
                "url": "corresponding url"
            }}

            Note: Don't choose "https://chatgpt.com", "https://chat.com", alternatives kind of domain as it requires signin.
        """
    )
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", api_key=GOOGLE_API_KEY).with_structured_output(TopicDetails)
    try:
        chain = prompt | model
        chosen_topic = chain.invoke({"topics": topics})
        return chosen_topic
    except Exception as e:
        print(f"❌ Auto choose failed: {e}")
        fallback = topics[0]
        return TopicDetails(title=fallback["title"], url=fallback["url"])

def fetch_blog_text(url: str) -> str:
    try:
        html = requests.get(url, timeout=10).text
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style"]):
            tag.extract()
        text = " ".join(soup.stripped_strings)
        return text[:15000]
    except Exception as e:
        return f"Error fetching {url}: {e}"

def summarize_blog(url: str) -> str:
    text = fetch_blog_text(url)
    prompt=f"Summarize this blog in 5 sentences:\n\n{text}"
    response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
    return response.text.strip()

def blend_multiple_blogs(summaries: List[str], topic_titles: List[str]) -> str:
    """Blend multiple blog summaries into a coherent research overview"""
    prompt = f"""
    You are an expert research analyst. Combine insights from multiple sources to create a comprehensive overview.
    
    Source Summaries:
    {chr(10).join([f"Source {i+1} ({topic_titles[i]}): {summary}" for i, summary in enumerate(summaries)])}
    
    Create a unified research summary that:
    1. Identifies common themes across all sources
    2. Highlights unique insights from each source
    3. Creates a coherent narrative flow
    4. Points out any contradictions or different perspectives
    
    Keep the summary comprehensive but concise (about 10-15 sentences).
    """
    response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
    return response.text.strip()

def generate_blog(topic: str, summary: str = "", tone: str = "professional", word_target: int = 700) -> BlogOutput:
    prompt = ChatPromptTemplate.from_template(
        """
        You are an expert content writer. Write clear, engaging, and original long-form blogs.

        Generate a long-form blog post ~{word_target} words about this topic: '{topic}'.
        context: {summary}
        Tone: {tone}.

        Requirements:
        - Provide output in JSON with keys: title, subtitle, body, seo_meta
        - Title: short, engaging (<=12 words)
        - Subtitle: catchy (<=20 words)
        - Body: split into multiple sections with headings
        - Tags: top five key for this blog
        - SEO Meta: concise description (<=160 chars)

        Format:
        {{
            "title": "short, engaging (<=12 words)",
            "subtitle": "catchy (<=20 words)",
            "body": [
                {{"heading": "Section Heading", "content": "Paragraph text"}},
                ...
            ],
            "conclusion": "typically makes up around 10% of the total document's word count, or about 5-7% for theses and dissertations.",
            "tags": ["keyword1", "keyword2"],
            "seo_meta": "meta description (<=160 chars)"
        }}

        Return valid JSON ONLY.
        """
        )
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", api_key=GOOGLE_API_KEY).with_structured_output(BlogOutput)
    try:
        chain = prompt | model
        print(f"📝 Generating blog for topic: {topic}")
        blog = chain.invoke({"topic": topic, "summary": summary, "tone": tone, "word_target": word_target})
        return blog
    except ValidationError as e:
        print("❌ Validation failed:", e)
        raise
    except Exception:
        return BlogOutput(
            title=topic,
            subtitle="",
            body=[{"heading": "Generated Content", "content": "Blog generation failed, please retry."}],
            conclusion="",
            tags=[],
            seo_meta=""
        )
    
def linkedIn_blog(blog) -> str:
    """
    Convert blog content to LinkedIn-optimized format using Unicode styling
    """
    prompt = f"""
    Convert this blog content into a LinkedIn-optimized post that will perform well on the platform.
    
    BLOG CONTENT:
    {blog}
    
    Requirements:
    1. Use Unicode characters for formatting (mathematical bold for emphasis, etc.)
    2. Keep it professional yet engaging
    3. Structure with clear sections and bullet points using Unicode symbols
    4. Include relevant hashtags at the end
    5. Optimize for readability and engagement
    6. Length: 800-1200 characters ideal for LinkedIn
    7. Start with a strong hook
    8. Use spacing and emojis strategically (1-2 emojis max)
    9. Include a call-to-action
    
    Formatting Guidelines:
    - Use mathematical bold script (𝗹𝗶𝗸𝗲 𝘁𝗵𝗶𝘀) for key points and headings
    - Use bullet points with • symbol
    - Use proper line breaks for readability
    
    Return ONLY the formatted LinkedIn post content, no explanations.
    """
    
    try:
        response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
        content = response.text.strip()
        
        # Ensure we have proper formatting
        if not any(char in content for char in ['𝗔', '𝘈', '•', '#']):
            # If no formatting detected, add basic structure
            lines = content.split('\n')
            formatted_lines = []
            for i, line in enumerate(lines):
                if i == 0:  # First line - make it bold
                    formatted_lines.append(f"𝗔{line[1:]}𝗔" if line.startswith('A') else f"𝗔{line}𝗔")
                elif line.strip() and not line.startswith('#'):
                    formatted_lines.append(f"• {line}")
                else:
                    formatted_lines.append(line)
            content = '\n'.join(formatted_lines)
        
        return content
        
    except Exception as e:
        print(f"❌ LinkedIn conversion failed: {e}")
        # Fallback basic formatting
        if isinstance(blog, str):
            return f"📝 {blog[:500]}... #Blog #Content"
        else:
            return "🚀 New insights on trending topics! Check out the full blog for more details. #ContentMarketing #Blogging"
        
def export_blog_to_docx(blog, filename):
    doc = Document()

    doc.add_heading(blog.title, 0)
    if blog.subtitle:
        doc.add_paragraph(blog.subtitle, style="Subtitle")

    doc.add_paragraph("")

    for section in blog.body:
        if section.heading:
            doc.add_heading(section.heading, 2)
        if section.content:
            doc.add_paragraph(section.content)

        doc.add_paragraph("")

    if hasattr(blog, "conclusion") and blog.conclusion:
        doc.add_heading("Conclusion", 2)
        doc.add_paragraph(blog.conclusion)
        doc.add_paragraph("")

    if blog.seo_meta:
        p = doc.add_paragraph(blog.seo_meta)
        run = p.runs[0]
        run.italic = True
        doc.add_paragraph("")
        
    if blog.tags:
        doc.add_heading("Tags", 2)
        for tag in blog.tags:
            doc.add_paragraph(tag, style="List Bullet")

    safe_filename = re.sub(r'[<>:"/\\|?*]', '_', blog.title)
    safe_filename = safe_filename.replace(" ", "_")
    filepath = blogs_folder / f"{safe_filename}.docx"

    doc.save(filepath)
    print(f"✅ Blog exported to {filepath}")


@app.get("/")
async def root():
    """Root endpoint that returns a hello message"""
    return {"message": "Oh you clever boy!", "status": "success", "service": "Blog Generator API"}

@app.get("/search-topics")
async def search_topics(query: str = Query(..., description="Search query for trending blogs")):
    """Search for topics and return results for user selection"""
    results = google_search(query)
    topics = [{"title": r.get("title"), "url": r.get("link")} for r in results]
    
    if not topics:
        raise HTTPException(status_code=404, detail="No search results found")
    
    # Store search results temporarily
    search_id = str(uuid.uuid4())
    search_store[search_id] = topics
    
    # Save search session to database
    save_search_session(search_id, query, topics)
    
    return {
        "search_id": search_id,
        "topics": topics,
        "count": len(topics)
    }

@app.post("/generate-blog-from-topic")
async def generate_blog_from_topic(
    search_id: str,
    selected_topic_index: int = Query(..., ge=0, description="Index of the selected topic"),
    tone: str = Query("professional", description="Tone of the blog"),
    word_target: int = Query(700, description="Target word count")
):
    """Generate blog from user-selected topic"""
    if search_id not in search_store:
        # Try to retrieve from database
        topics = get_search_session(search_id)
        if not topics:
            raise HTTPException(status_code=404, detail="Search session not found")
        search_store[search_id] = topics
    else:
        topics = search_store[search_id]
    
    if selected_topic_index >= len(topics):
        raise HTTPException(status_code=400, detail="Invalid topic index")
    
    chosen_topic = topics[selected_topic_index]
    chosen_item = TopicDetails(title=chosen_topic["title"], url=chosen_topic["url"])
    
    summary = summarize_blog(chosen_item.url)
    blog = generate_blog(chosen_item.title, summary, tone, word_target)
    linkedIn_version = linkedIn_blog(blog)
    
    safe_filename = blog.title.replace(" ", "_").replace("/", "_")
    filepath = f"{safe_filename}.docx"
    export_blog_to_docx(blog, safe_filename)
    
    blog_dict = blog.model_dump()
    
    # Save to database
    blog_id = save_blog_outcome(
        search_id=search_id,
        title=blog.title,
        subtitle=blog.subtitle,
        content_json=blog_dict,
        summary=summary,
        linkedin_version=linkedIn_version,
        linkedin_content=linkedIn_version,  # Store the actual content
        docx_filename=filepath,
        tone=tone,
        word_count=word_target,
        tags=blog.tags,
        seo_meta=blog.seo_meta
    )
    
    # Also save to LinkedIn posts table
    save_linkedin_post(blog_id, 'regular', linkedIn_version)
    
    # Clean up stored search results
    if search_id in search_store:
        del search_store[search_id]
    
    return {
        "chosen_topic": chosen_item,
        "summary": summary,
        "docx_file": filepath,
        "blog_content": blog_dict,
        "linkedIn_version": linkedIn_version,
        "blog_id": blog_id
    }

@app.post("/generate-blended-blog")
async def generate_blended_blog(request: MultiBlogRequest):
    """Generate a blog by blending multiple selected topics"""
    print("------1")
    
    # Try to get from database if not in memory
    if request.search_id not in search_store:
        topics = get_search_session(request.search_id)
        if not topics:
            raise HTTPException(status_code=404, detail="Search session not found")
        search_store[request.search_id] = topics
    
    topics = search_store[request.search_id]
    
    print("------3")
    # Validate all topic indices
    for index in request.topic_indices:
        print("------4")
        if index >= len(topics):
            print("------5")
            raise HTTPException(status_code=400, detail=f"Invalid topic index: {index}")
    
    print("------6")
    # Get selected topics
    selected_topics = [topics[i] for i in request.topic_indices]
    
    print("------7")
    # Summarize each blog
    summaries = []
    print("------8")
    for topic in selected_topics:
        print("------9")
        summary = summarize_blog(topic["url"])
        print("------10")
        summaries.append(summary)
    
    print("------11")
    # Blend summaries
    blended_summary = blend_multiple_blogs(
        summaries, 
        [topic["title"] for topic in selected_topics]
    )
    print("------12")
    
    # Create a combined topic title
    combined_title = f"Comprehensive Analysis: {', '.join([topic['title'][:30] for topic in selected_topics[:3]])}"
    print("------13")
    if len(selected_topics) > 3:
        print("------14")
        combined_title += " and more"
    
    # Generate blog from blended content
    blog = generate_blog(combined_title, blended_summary, request.tone, request.word_target)
    linkedIn_version = linkedIn_blog(blog)
    
    safe_filename = f"blended_blog_{uuid.uuid4().hex[:8]}.docx"
    filepath = safe_filename
    export_blog_to_docx(blog, safe_filename)
    
    blog_dict = blog.model_dump()
    
    # Save to database
    blog_id = save_blended_blog_outcome(
        search_id=request.search_id,
        title=blog.title,
        subtitle=blog.subtitle,
        content_json=blog_dict,
        blended_summary=blended_summary,
        linkedin_version=linkedIn_version,
        linkedin_content=linkedIn_version,  # Store the actual content
        docx_filename=filepath,
        topic_indices=request.topic_indices,
        individual_summaries=summaries,
        tone=request.tone,
        word_count=request.word_target,
        tags=blog.tags,
        seo_meta=blog.seo_meta
    )
    
    # Also save to LinkedIn posts table
    save_linkedin_post(blog_id, 'blended', linkedIn_version)
    
    # Clean up stored search results
    if request.search_id in search_store:
        del search_store[request.search_id]
    
    return {
        "selected_topics": selected_topics,
        "blended_summary": blended_summary,
        "docx_file": filepath,
        "blog_content": blog_dict,
        "linkedIn_version": linkedIn_version,
        "individual_summaries": summaries,
        "blog_id": blog_id
    }

# Keep original endpoint for backward compatibility
@app.get("/generate-blog")
def generate_blog_endpoint(
    query: str = Query(..., description="Search query for trending blogs"), 
    tone: str = Query("professional", description="Tone of the blog"),
    word_target: int = Query(700, description="Target word count")
):
    results = google_search(query)
    topics = [{"title": r.get("title"), "url": r.get("link")} for r in results]
    if not topics:
        return {"error": "No search results found"}

    chosen_item = choose_topic_auto(topics)
    if not chosen_item or not chosen_item.title:
        return {"error": "Gemini did not return a valid topic"}

    summary = summarize_blog(chosen_item.url)
    print("==============summary", summary)
    blog = generate_blog(chosen_item.title, summary, tone, word_target)
    linkedIn_version = linkedIn_blog(blog)
    safe_filename = blog.title.replace(" ", "_").replace("/", "_")
    filepath = f"{safe_filename}.docx"
    export_blog_to_docx(blog, safe_filename)
    blog_dict = blog.model_dump()
    
    # Save to database (generate a search_id for this legacy endpoint)
    search_id = f"legacy_{uuid.uuid4().hex[:8]}"
    blog_id = save_blog_outcome(
        search_id=search_id,
        title=blog.title,
        subtitle=blog.subtitle,
        content_json=blog_dict,
        summary=summary,
        linkedin_version=linkedIn_version,
        linkedin_content=linkedIn_version,
        docx_filename=filepath,
        tone=tone,
        word_count=word_target,
        tags=blog.tags,
        seo_meta=blog.seo_meta
    )
    
    # Also save to LinkedIn posts table
    save_linkedin_post(blog_id, 'regular', linkedIn_version)
    
    return {
        "topics": topics,
        "chosen_topic": chosen_item,
        "summary": summary,
        "docx_file": filepath,
        "blog_content": blog_dict,
        "linkedIn_version": linkedIn_version,
        "blog_id": blog_id
    }

@app.post("/suggest-headings")
def suggest_headings(request: dict):
    """
    Suggest multiple possible section headings for a blog body.
    """
    from langchain_google_genai import ChatGoogleGenerativeAI

    body_content = request.get("body_content", "")
    topic = request.get("topic", "")
    if not body_content:
        raise HTTPException(status_code=400, detail="Body content required")

    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash")
    prompt = f"""
    Suggest 5 possible section headings for this blog content.
    Topic: {topic}

    Content:
    {body_content[:1000]}

    Return a JSON array of short, descriptive section headings.
    """
    response = llm.invoke([{"role": "user", "content": prompt}])
    try:
        headings = json.loads(response.content[0].text)
    except Exception:
        headings = [h.strip() for h in response.text.split("\n") if h.strip()][:5]
    return {"headings": headings}


@app.post("/generate-auto-blog")
async def generate_auto_blog(
    search_id: str,
    tone: str = "professional",
    word_target: int = 700
):
    """Automatically choose the best topic and generate blog"""
    if search_id not in search_store:
        # Try to retrieve from database
        topics = get_search_session(search_id)
        if not topics:
            raise HTTPException(status_code=404, detail="Search session not found")
        search_store[search_id] = topics
    
    topics = search_store[search_id]
    
    # Use the existing auto-topic selection logic
    chosen_item = choose_topic_auto(topics)
    
    # Generate summary and blog
    summary = summarize_blog(chosen_item.url)
    blog = generate_blog(chosen_item.title, summary, tone, word_target)
    linkedIn_version = linkedIn_blog(blog)
    
    safe_filename = blog.title.replace(" ", "_").replace("/", "_")
    filepath = f"{safe_filename}.docx"
    export_blog_to_docx(blog, safe_filename)
    
    blog_dict = blog.model_dump()
    
    # Save to database
    blog_id = save_blog_outcome(
        search_id=search_id,
        title=blog.title,
        subtitle=blog.subtitle,
        content_json=blog_dict,
        summary=summary,
        linkedin_version=linkedIn_version,
        linkedin_content=linkedIn_version,
        docx_filename=filepath,
        tone=tone,
        word_count=word_target,
        tags=blog.tags,
        seo_meta=blog.seo_meta
    )
    
    # Also save to LinkedIn posts table
    save_linkedin_post(blog_id, 'regular', linkedIn_version)
    
    # Clean up stored search results
    if search_id in search_store:
        del search_store[search_id]
    
    return {
        "chosen_topic": chosen_item,
        "summary": summary,
        "docx_file": filepath,
        "blog_content": blog_dict,
        "linkedIn_version": linkedIn_version,
        "auto_selected": True,
        "blog_id": blog_id
    }

@app.get("/blogs")
def list_blogs():
    """List all generated blogs from database"""
    blog_files = [f.name for f in blogs_folder.glob("*.docx")]
    
    # Get blogs from database
    regular_blogs = get_all_blogs()
    blended_blogs = get_blended_blogs()
    
    return {
        "docx_files": blog_files,
        "regular_blogs": regular_blogs,
        "blended_blogs": blended_blogs,
        "total_regular": len(regular_blogs),
        "total_blended": len(blended_blogs)
    }

@app.get("/blogs/{blog_id}")
async def get_blog(blog_id: int, blog_type: str = Query("regular", description="Type of blog: regular or blended")):
    """Get specific blog by ID"""
    if blog_type == "regular":
        blog = get_blog_by_id(blog_id)
    elif blog_type == "blended":
        blog = get_blended_blog_by_id(blog_id)
    else:
        raise HTTPException(status_code=400, detail="Invalid blog type. Use 'regular' or 'blended'")
    
    if not blog:
        raise HTTPException(status_code=404, detail="Blog not found")
    
    return blog

@app.get("/blogs/search/{search_id}")
async def get_blogs_by_search(search_id: str):
    """Get all blogs generated from a specific search session"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Get regular blogs
    cursor.execute('SELECT * FROM blogs WHERE search_id = ? ORDER BY created_at DESC', (search_id,))
    regular_blogs = []
    for row in cursor.fetchall():
        # Get Instagram posts for each blog
        cursor.execute('SELECT carousel_text, images_json FROM instagram_posts WHERE blog_id = ? AND blog_type = "regular"', (row[0],))
        insta_data = cursor.fetchone()
        
        # Get LinkedIn posts for each blog
        cursor.execute('SELECT post_content FROM linkedin_posts WHERE blog_id = ? AND blog_type = "regular"', (row[0],))
        linkedin_data = cursor.fetchone()
        
        regular_blogs.append({
            'id': row[0],
            'search_id': row[1],
            'title': row[2],
            'subtitle': row[3],
            'content_json': json.loads(row[4]),
            'summary': row[5],
            'linkedin_version': row[6],
            'linkedin_content': row[7],
            'docx_filename': row[8],
            'tone': row[9],
            'word_count': row[10],
            'tags': json.loads(row[11]) if row[11] else [],
            'seo_meta': row[12],
            'created_at': row[13],
            'updated_at': row[14],
            'instagram_carousel': {
                'carousel_text': insta_data[0] if insta_data else None,
                'images': json.loads(insta_data[1]) if insta_data and insta_data[1] else []
            } if insta_data else None,
            'linkedin_post': linkedin_data[0] if linkedin_data else None
        })
    
    # Get blended blogs
    cursor.execute('SELECT * FROM blended_blogs WHERE search_id = ? ORDER BY created_at DESC', (search_id,))
    blended_blogs = []
    for row in cursor.fetchall():
        # Get Instagram posts for each blended blog
        cursor.execute('SELECT carousel_text, images_json FROM instagram_posts WHERE blog_id = ? AND blog_type = "blended"', (row[0],))
        insta_data = cursor.fetchone()
        
        # Get LinkedIn posts for each blended blog
        cursor.execute('SELECT post_content FROM linkedin_posts WHERE blog_id = ? AND blog_type = "blended"', (row[0],))
        linkedin_data = cursor.fetchone()
        
        blended_blogs.append({
            'id': row[0],
            'search_id': row[1],
            'title': row[2],
            'subtitle': row[3],
            'content_json': json.loads(row[4]),
            'blended_summary': row[5],
            'linkedin_version': row[6],
            'linkedin_content': row[7],
            'docx_filename': row[8],
            'topic_indices': json.loads(row[9]) if row[9] else [],
            'individual_summaries': json.loads(row[10]) if row[10] else [],
            'tone': row[11],
            'word_count': row[12],
            'tags': json.loads(row[13]) if row[13] else [],
            'seo_meta': row[14],
            'created_at': row[15],
            'updated_at': row[16],
            'instagram_carousel': {
                'carousel_text': insta_data[0] if insta_data else None,
                'images': json.loads(insta_data[1]) if insta_data and insta_data[1] else []
            } if insta_data else None,
            'linkedin_post': linkedin_data[0] if linkedin_data else None
        })
    
    conn.close()
    
    return {
        "search_id": search_id,
        "regular_blogs": regular_blogs,
        "blended_blogs": blended_blogs,
        "total_regular": len(regular_blogs),
        "total_blended": len(blended_blogs)
    }

@app.post("/generate-instagram-carousel")
async def generate_instagram_carousel_endpoint(request: InstagramRequest):
    """Generate Instagram carousel for a blog and save to database"""
    print(request.blog_content, request.blog_id, request.blog_type, request.customization_prompt)
    try:
        print(f"📱 Generating Instagram carousel for blog content...")
        
        # Step 1: Generate carousel text
        carousel_text = generate_instagram_carousel(request.blog_content)
        print(f"✅ Carousel text generated: {len(carousel_text)} characters")
        
        # Step 2: Extract slides
        slides = extract_slides(carousel_text)
        print(f"✅ Extracted {len(slides)} slides")
        
        # Step 3: Generate images for each slide as base64
        image_filenames = {}  # store {filename: base64_data}
        
        for i, (title, content) in enumerate(slides.items(), start=1):
            try:
                print(f"🖼️ Generating image for slide {i}: {title[:50]}...")
                base64_data = generate_slide_image(content.strip(), i, request.customization_prompt)
                
                # Generate filename key
                filename = f"slide_{i}_{uuid.uuid4().hex[:8]}.png"
                
                # Store filename: base64_data mapping
                image_filenames[filename] = base64_data
                
                print(f"✅ Slide {i} image generated and stored")
            except Exception as e:
                print(f"❌ Error generating slide {i} image: {e}")
                # Add a placeholder entry
                placeholder_name = f"slide_{i}_error_placeholder.png"
                image_filenames[placeholder_name] = None
        
        # Step 4: Save to database if blog_id is provided
        if request.blog_id:
            try:
                # Ensure DB column type supports JSON
                # If not, modify the schema to JSONB / TEXT with JSON serialization
                save_instagram_post(
                    request.blog_id, 
                    request.blog_type,
                    carousel_text, 
                    image_filenames  # dict {filename: base64_data}
                )
                print(f"💾 Saved to database for blog_id: {request.blog_id}")
            except Exception as e:
                print(f"⚠️ Could not save to database: {e}")
        
        return {
            "carousel_text": carousel_text,
            "image_filenames": image_filenames,  # dict {filename: base64}
            "slides_count": len(slides),
            "blog_id": request.blog_id,
            "blog_type": request.blog_type,
            "customization": request.customization_prompt,
            "status": "success"
        }
        
    except Exception as e:
        print(f"❌ Carousel generation failed: {e}")
        raise HTTPException(status_code=500, detail=f"Carousel generation failed: {e}")


@app.get("/instagram-posts/{blog_id}")
async def get_instagram_posts(blog_id: int, blog_type: str = Query("regular")):
    """Get Instagram carousel posts for a specific blog"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT carousel_text, images_json, created_at 
        FROM instagram_posts 
        WHERE blog_id = ? AND blog_type = ?
        ORDER BY created_at DESC
    ''', (blog_id, blog_type))
    
    posts = []
    for row in cursor.fetchall():
        posts.append({
            'carousel_text': row[0],
            'images': json.loads(row[1]) if row[1] else [],
            'created_at': row[2]
        })
    
    conn.close()
    
    return {
        "blog_id": blog_id,
        "blog_type": blog_type,
        "instagram_posts": posts
    }

@app.get("/linkedin-posts/{blog_id}")
async def get_linkedin_posts(blog_id: int, blog_type: str = Query("regular")):
    """Get LinkedIn posts for a specific blog"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT post_content, hashtags, character_count, engagement_score, created_at 
        FROM linkedin_posts 
        WHERE blog_id = ? AND blog_type = ?
        ORDER BY created_at DESC
    ''', (blog_id, blog_type))
    
    posts = []
    for row in cursor.fetchall():
        posts.append({
            'post_content': row[0],
            'hashtags': row[1],
            'character_count': row[2],
            'engagement_score': row[3],
            'created_at': row[4]
        })
    
    conn.close()
    
    return {
        "blog_id": blog_id,
        "blog_type": blog_type,
        "linkedin_posts": posts
    }

@app.post("/refresh-section")
async def refresh_section(request: dict):
    """
    Refresh or generate content for a new section heading using AI, 
    considering the context of the existing blog.
    """
    try:
        title = request.get("title", "")
        existing_body = request.get("existing_body", [])
        new_heading = request.get("new_heading", "")

        if not new_heading:
            raise HTTPException(status_code=400, detail="Missing 'new_heading' in request")

        # Build contextual prompt
        context_text = "\n".join(
            [f"{sec['heading']}: {sec.get('content', '')}" for sec in existing_body if sec.get('heading')]
        )
        prompt = f"""
        You are an expert content writer. The blog is titled: "{title}".

        Existing sections:
        {context_text}

        A new section has been added with heading: "{new_heading}".
        Write 2-3 engaging paragraphs for this new section that fit naturally with the existing content.

        Keep the tone consistent and ensure the new content flows well with what's already written.
        Return only the content paragraphs, no heading or markdown.
        """

        response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
        new_content = response.text.strip()

        return {
            "new_heading": new_heading,
            "new_content": new_content,
            "status": "success"
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate section: {str(e)}")

@app.post("/refresh-heading")
async def refresh_heading(request: dict):
    """
    Generate 2-3 alternative heading options for a given section.
    """
    try:
        title = request.get("title", "")
        existing_body = request.get("existing_body", [])
        section_index = request.get("section_index", 0)

        if section_index >= len(existing_body):
            raise HTTPException(status_code=400, detail="Invalid section index")

        current_heading = existing_body[section_index].get("heading", "")
        current_content = existing_body[section_index].get("content", "")

        prompt = f"""
        This is a blog titled: "{title}".

        Current heading: "{current_heading}"
        Content: "{current_content}"

        Generate 3 alternative heading options for this section.
        Make them engaging, SEO-friendly, and fitting for the content.
        Return as a JSON array: ["Option 1", "Option 2", "Option 3"]
        """

        response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
        # Extract JSON array from response
        import re
        match = re.search(r'\[.*\]', response.text)
        if match:
            import json
            headings = json.loads(match.group())
        else:
            # Fallback: split by lines and clean
            headings = [line.strip(' -•"') for line in response.text.split('\n') if line.strip()]
            headings = [h for h in headings if h][:3]

        return {
            "current_heading": current_heading,
            "alternatives": headings,
            "status": "success"
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate headings: {str(e)}")

@app.post("/generate-with-ai")
async def generate_with_ai(request: dict):
    """
    Generate content for a given heading using AI.
    """
    try:
        title = request.get("title", "")
        existing_body = request.get("existing_body", [])
        heading = request.get("heading", "")

        if not heading:
            raise HTTPException(status_code=400, detail="Missing 'heading' in request")

        # Build contextual prompt
        context_text = "\n".join(
            [f"{sec['heading']}: {sec.get('content', '')}" for sec in existing_body if sec.get('heading')]
        )
        prompt = f"""
        You are an expert content writer. The blog is titled: "{title}".

        Existing sections:
        {context_text}

        Write content for the section: "{heading}".
        Generate 2-3 engaging paragraphs that fit naturally with the existing content.

        Keep the tone consistent and ensure the content flows well with what's already written.
        Return only the content paragraphs, no heading or markdown.
        """

        response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
        content = response.text.strip()

        return {
            "heading": heading,
            "content": content,
            "status": "success"
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate content: {str(e)}")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)