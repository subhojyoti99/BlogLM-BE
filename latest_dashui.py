import streamlit as st
import requests
import json
import os
from datetime import datetime
import base64

# Streamlit app configuration
st.set_page_config(
    page_title="Trending Blog Generator",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Session state initialization
if "search_results" not in st.session_state:
    st.session_state.search_results = None
if "search_id" not in st.session_state:
    st.session_state.search_id = None
if "selected_topics" not in st.session_state:
    st.session_state.selected_topics = []
if "generation_mode" not in st.session_state:
    st.session_state.generation_mode = "single"  # or "blended"

# Custom CSS styling
st.markdown("""
    <style>
    .main-header {
        font-size: 3rem;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .blog-card {
        background-color: #f8f9fa;
        padding: 1.5rem;
        border-radius: 10px;
        border-left: 5px solid #1f77b4;
        margin-bottom: 1rem;
    }
    .topic-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #ddd;
        margin-bottom: 0.5rem;
        cursor: pointer;
    }
    .topic-card.selected {
        border: 2px solid #1f77b4;
        background-color: #e3f2fd;
    }
    .section-header {
        color: #1f77b4;
        border-bottom: 2px solid #1f77b4;
        padding-bottom: 0.5rem;
        margin-top: 1.5rem;
    }
    .stButton button {
        background-color: #1f77b4;
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 5px;
    }
    .stButton button:hover {
        background-color: #0d6efd;
    }
    </style>
""", unsafe_allow_html=True)

class BlogGeneratorUI:
    def __init__(self):
        self.api_url = "http://localhost:8000"
        
    def create_download_link(self, file_path):
        """Create a download link for the generated DOCX file"""
        try:
            with open(file_path, "rb") as file:
                data = file.read()
            b64 = base64.b64encode(data).decode()
            filename = os.path.basename(file_path)
            href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">📥 Download DOCX</a>'
            return href
        except FileNotFoundError:
            return "File not found yet. Please wait for generation to complete."
    
    def search_topics(self, query):
        """Search for topics using the API"""
        try:
            response = requests.get(
                f"{self.api_url}/search-topics",
                params={"query": query},
                timeout=30
            )
            if response.status_code == 200:
                return response.json()
            else:
                st.error(f"Search failed: {response.status_code}")
                return None
        except Exception as e:
            st.error(f"Search error: {str(e)}")
            return None
    
    def generate_single_blog(self, search_id, topic_index, tone, word_target):
        """Generate blog from single topic"""
        try:
            response = requests.post(
                f"{self.api_url}/generate-blog-from-topic",
                params={
                    "search_id": search_id,
                    "selected_topic_index": topic_index,
                    "tone": tone,
                    "word_target": word_target
                },
                timeout=120
            )
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"Generation failed: {response.status_code}"}
        except Exception as e:
            return {"error": f"Generation error: {str(e)}"}
    
    def generate_blended_blog(self, search_id, topic_indices, tone, word_target):
        """Generate blog by blending multiple topics"""
        try:
            response = requests.post(
                f"{self.api_url}/generate-blended-blog",
                json={
                    "search_id": search_id,
                    "topic_indices": topic_indices,
                    "tone": tone,
                    "word_target": word_target
                },
                timeout=180
            )
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"Blending failed: {response.status_code}"}
        except Exception as e:
            return {"error": f"Blending error: {str(e)}"}
    
    def display_topic_selection(self, topics, search_id):
        """Display topics for user selection"""
        st.markdown("### 🔍 Found Topics")
        st.info(f"Found {len(topics)} topics. Select one or multiple topics to generate a blog.")
        
        # Generation mode selection
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🎯 Single Topic Mode", use_container_width=True):
                st.session_state.generation_mode = "single"
                st.session_state.selected_topics = []
        with col2:
            if st.button("🔄 Blended Topics Mode", use_container_width=True):
                st.session_state.generation_mode = "blended"
        
        st.markdown(f"**Current Mode:** {'🎯 Single Topic' if st.session_state.generation_mode == 'single' else '🔄 Blended Topics'}")
        
        # Display topics with selection
        selected_indices = []
        for i, topic in enumerate(topics):
            col1, col2 = st.columns([1, 20])
            with col1:
                if st.session_state.generation_mode == "single":
                    selected = st.radio(f"Select", [f"Topic {i+1}"], key=f"topic_{i}", label_visibility="collapsed")
                    if selected:
                        selected_indices = [i]
                else:
                    if st.checkbox("", key=f"check_{i}"):
                        selected_indices.append(i)
            with col2:
                with st.container():
                    st.markdown(f"**{topic['title']}**")
                    st.caption(f"🔗 {topic['url']}")
        
        return selected_indices
    
    def display_blog_content(self, blog_data, is_blended=False):
        """Display and allow editing of generated blog content"""
        if "error" in blog_data:
            st.error(f"Error: {blog_data['error']}")
            return

        if is_blended:
            st.markdown("### 🎉 Blended Blog Generated!")
            st.markdown("#### Selected Sources:")
            for i, topic in enumerate(blog_data['selected_topics']):
                st.write(f"{i+1}. **{topic['title']}**")
        else:
            st.markdown(f"### 📖 Generated Blog: {blog_data['chosen_topic']['title']}")

        # Editable blog content
        if 'blog_content' in blog_data:
            blog = blog_data['blog_content']

            st.subheader("✏️ Edit Your Blog")
            
            col1, col2 = st.columns(2)
            with col1:
                edited_title = st.text_input("Blog Title", blog['title'])
            with col2:
                edited_subtitle = st.text_area("Subtitle", blog['subtitle'])

            edited_sections = []
            for i, section in enumerate(blog['body']):
                st.markdown(f"#### Section {i+1}")
                col1, col2 = st.columns([1, 3])
                with col1:
                    heading = st.text_input(f"Heading {i+1}", section['heading'], key=f"heading_{i}")
                with col2:
                    content = st.text_area(f"Content {i+1}", section['content'], height=150, key=f"content_{i}")
                edited_sections.append({"heading": heading, "content": content})

            edited_conclusion = st.text_area("Conclusion", blog.get("conclusion", ""))
            edited_tags = st.text_input("Tags (comma separated)", ", ".join(blog.get("tags", [])))
            edited_seo = st.text_area("SEO Meta Description", blog.get("seo_meta", ""))

            if st.button("📥 Export as DOCX"):
                from docx import Document
                doc = Document()

                doc.add_heading(edited_title, 0)
                if edited_subtitle:
                    doc.add_paragraph(edited_subtitle, style="Intense Quote")

                for section in edited_sections:
                    doc.add_heading(section["heading"], level=1)
                    doc.add_paragraph(section["content"])

                if edited_conclusion:
                    doc.add_heading("Conclusion", level=1)
                    doc.add_paragraph(edited_conclusion)

                if edited_tags:
                    doc.add_heading("Tags", level=2)
                    doc.add_paragraph(edited_tags)

                if edited_seo:
                    p = doc.add_paragraph(edited_seo)
                    if p.runs:
                        p.runs[0].italic = True

                file_path = f"blog_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                doc.save(file_path)
                st.success("✅ Blog exported successfully!")
                st.markdown(self.create_download_link(file_path), unsafe_allow_html=True)

            # Show LinkedIn version
            with st.expander("📱 LinkedIn Version"):
                st.write(blog_data.get('linkedIn_version', ''))

            if is_blended:
                with st.expander("📊 Individual Summaries"):
                    for i, summary in enumerate(blog_data.get('individual_summaries', [])):
                        st.write(f"**Source {i+1}:** {summary}")
    
    def render_sidebar(self):
        """Render the sidebar with input controls"""
        with st.sidebar:
            st.markdown("## ⚙️ Blog Generation Settings")
            
            # Search query input
            query = st.text_input(
                "🔍 Search Query",
                value="latest technology trends 2025",
                help="Enter a search term to find trending topics"
            )
            
            # Tone selection
            tone = st.selectbox(
                "🎭 Writing Tone",
                options=["professional", "casual", "persuasive", "informative", "enthusiastic"],
                index=0,
                help="Select the tone for your blog post"
            )
            
            # Word target input
            word_target = st.slider(
                "📏 Target Word Count",
                min_value=300,
                max_value=2000,
                value=700,
                step=100,
                help="Approximate word count for the generated blog"
            )
            
            # Search button
            search_btn = st.button(
                "🔍 Search Topics",
                type="primary",
                use_container_width=True
            )
            
            st.markdown("---")
            st.markdown("### 📊 Current Session")
            if st.session_state.search_results:
                st.success(f"✅ {len(st.session_state.search_results)} topics found")
                st.info(f"Mode: {'🎯 Single' if st.session_state.generation_mode == 'single' else '🔄 Blended'}")
            else:
                st.info("No active search session")
            
            st.markdown("---")
            st.markdown("### ℹ️ About")
            st.caption("""
                **New Features:**
               - 🎯 Select individual topics
                - 🔄 Blend multiple topics
                - ✏️ Edit generated content
                - 📥 Export to DOCX
            """)
            
            return query, tone, word_target, search_btn
    
    def render_main_content(self):
        """Render the main content area"""
        st.markdown('<h1 class="main-header">📝 AI Blog Generator</h1>', unsafe_allow_html=True)
        st.markdown("Generate engaging blog posts from trending topics using AI - **Now with topic selection and blending!**")

        # Sidebar inputs
        query, tone, word_target, search_btn = self.render_sidebar()

        # Main content
        if search_btn:
            if not query.strip():
                st.error("Please enter a search query")
                return

            with st.spinner("🔍 Searching for trending topics..."):
                result = self.search_topics(query)
                if result:
                    st.session_state.search_results = result['topics']
                    st.session_state.search_id = result['search_id']
                    st.rerun()

        if st.session_state.search_results:
            selected_indices = self.display_topic_selection(
                st.session_state.search_results, 
                st.session_state.search_id
            )
            
            if selected_indices:
                if st.session_state.generation_mode == "single" and len(selected_indices) == 1:
                    if st.button("🚀 Generate Single Blog", type="primary", use_container_width=True):
                        with st.spinner("Generating your blog post..."):
                            result = self.generate_single_blog(
                                st.session_state.search_id,
                                selected_indices[0],
                                tone,
                                word_target
                            )
                            self.display_blog_content(result, is_blended=False)
                
                elif st.session_state.generation_mode == "blended" and len(selected_indices) >= 2:
                    st.info(f"Selected {len(selected_indices)} topics for blending")
                    if st.button("🔄 Generate Blended Blog", type="primary", use_container_width=True):
                        with st.spinner("Blending multiple sources and generating comprehensive blog..."):
                            result = self.generate_blended_blog(
                                st.session_state.search_id,
                                selected_indices,
                                tone,
                                word_target
                            )
                            self.display_blog_content(result, is_blended=True)
                
                elif st.session_state.generation_mode == "blended" and len(selected_indices) < 2:
                    st.warning("Please select at least 2 topics for blended mode")
                
                else:
                    st.info("Please select a topic to generate a blog")

        else:
            # Show examples and guide when no search is active
            tab1, tab2 = st.tabs(["🚀 Get Started", "📚 Examples"])
            
            with tab1:
                st.markdown("""
                ### How to use the enhanced blog generator:
                
                1. **Enter a search query** in the sidebar
                2. **Click 'Search Topics'** to find trending content
                3. **Choose your generation mode:**
                   - 🎯 **Single Topic**: Pick one topic for focused content
                   - 🔄 **Blended Topics**: Select multiple topics for comprehensive analysis
                4. **Select topics** from the results
                5. **Generate and edit** your blog
                6. **Export** as DOCX
                """)
            
            with tab2:
                st.markdown("### 💡 Example Use Cases")
                
                examples = [
                    {
                        "query": "AI in healthcare machine learning",
                        "mode": "blended",
                        "description": "Get comprehensive coverage of AI applications across different medical fields"
                    },
                    {
                        "query": "sustainable fashion eco-friendly",
                        "mode": "single", 
                        "description": "Deep dive into one specific aspect of sustainable fashion"
                    },
                    {
                        "query": "remote work productivity tools",
                        "mode": "blended",
                        "description": "Combine multiple perspectives on remote work effectiveness"
                    }
                ]
                
                for example in examples:
                    with st.expander(f"🔍 {example['query']} ({example['mode'].title()} Mode)"):
                        st.write(example['description'])
                        if st.button("Try this example", key=example['query']):
                            st.session_state.query = example['query']
                            st.rerun()
    
    def run(self):
        """Run the Streamlit application"""
        try:
            # Check API connection
            try:
                response = requests.get(f"{self.api_url}/", timeout=2)
                if response.status_code != 200:
                    st.warning("⚠️ FastAPI server might not be running. Please start it with: `uvicorn blogger:app --reload`")
            except:
                st.warning("⚠️ Cannot connect to FastAPI server. Please make sure it's running on http://localhost:8000")
            
            self.render_main_content()
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

# Run the application
if __name__ == "__main__":
    app = BlogGeneratorUI()
    app.run()