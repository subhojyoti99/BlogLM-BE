# Human in the loop
# Multiple blogs blending

# Add ai content after adding new content box [human will add the heading rest if the content use button(generate with ai)], heading refresh with 2-3 options
# AI Generated content checker

# Image multiple with any white/black background check insta for reference
# Image Generation with category/style

# Image with quote
# Image with Film quote with relatable film image(quote on the generate/webscrap: for image)

# Custom image bg, suggestions new columns

import os
import re
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
import google.generativeai as genai
from typing import List, Dict
from pydantic import BaseModel, Field
import json
import uuid
from pydantic import ValidationError
from docx import Document
from fastapi import FastAPI, Query, HTTPException
import uvicorn
from fastapi.middleware.cors import CORSMiddleware
import asyncio
from pathlib import Path
from instaImages import generate_instagram_carousel, extract_slides, generate_slide_image

load_dotenv()

app = FastAPI(title="Trending Blog Generator API")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # For development only, restrict in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

blogs_folder = Path("blogs")
blogs_folder.mkdir(exist_ok=True)

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
GOOGLE_CX = os.getenv("GOOGLE_CX")
genai.configure(api_key=GOOGLE_API_KEY)

class BlogSection(BaseModel):
    heading: str = Field(..., description="Heading of the blog section")
    content: str = Field(..., description="Content/paragraph text for this section")

class BlogOutput(BaseModel):
    title: str = Field(..., description="Short, engaging blog title (<=12 words)")
    subtitle: str = Field(..., description="Catchy subtitle (<=20 words)")
    body: List[BlogSection] = Field(default_factory=list, description="List of blog sections with headings and content")
    conclusion: str = Field("", description="Closing summary paragraph (~10% of length)")
    tags: List[str] = Field(default_factory=list, description="Top 5 keyword tags for the blog")
    seo_meta: str = Field("", description="SEO meta description (<=160 characters)")

class TopicDetails(BaseModel):
    title: str = Field(..., description="Chosen blog title")
    url: str = Field(..., description="Corresponding blog URL")

class MultiBlogRequest(BaseModel):
    topic_indices: List[int] = Field(..., description="List of topic indices to blend")
    search_id: str = Field(..., description="Search session ID")
    tone: str = Field("professional", description="Tone of the blog")
    word_target: int = Field(700, description="Target word count")

search_store: Dict[str, List[Dict[str, str]]] = {}


def google_search(query, num_results=10):
    """Search Google with Custom Search API"""
    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "q": query,
        "key": GOOGLE_API_KEY,
        "cx": GOOGLE_CX,
        "num": num_results,
    }
    res = requests.get(url, params=params)
    res.raise_for_status()
    return res.json().get("items", [])

def choose_topic_auto(topics: List[Dict[str, str]]) -> TopicDetails:
    prompt = ChatPromptTemplate.from_template(
        """
            Here are some trending blog titles and links:
            {topics}

            Please select the single most trending and relevant blog.
            Return the result strictly in JSON format like this:
            {{
                "title": "chosen blog title",
                "url": "corresponding url"
            }}

            Note: Don't choose "https://chatgpt.com", "https://chat.com", alternatives kind of domain as it requires signin.
        """
    )
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", api_key=GOOGLE_API_KEY).with_structured_output(TopicDetails)
    try:
        chain = prompt | model
        chosen_topic = chain.invoke({"topics": topics})
        return chosen_topic
    except Exception as e:
        print(f"❌ Auto choose failed: {e}")
        fallback = topics[0]
        return TopicDetails(title=fallback["title"], url=fallback["url"])

def fetch_blog_text(url: str) -> str:
    try:
        html = requests.get(url, timeout=10).text
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style"]):
            tag.extract()
        text = " ".join(soup.stripped_strings)
        return text[:15000]
    except Exception as e:
        return f"Error fetching {url}: {e}"

def summarize_blog(url: str) -> str:
    text = fetch_blog_text(url)
    prompt=f"Summarize this blog in 5 sentences:\n\n{text}"
    response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
    return response.text.strip()

def blend_multiple_blogs(summaries: List[str], topic_titles: List[str]) -> str:
    """Blend multiple blog summaries into a coherent research overview"""
    prompt = f"""
    You are an expert research analyst. Combine insights from multiple sources to create a comprehensive overview.
    
    Source Summaries:
    {chr(10).join([f"Source {i+1} ({topic_titles[i]}): {summary}" for i, summary in enumerate(summaries)])}
    
    Create a unified research summary that:
    1. Identifies common themes across all sources
    2. Highlights unique insights from each source
    3. Creates a coherent narrative flow
    4. Points out any contradictions or different perspectives
    
    Keep the summary comprehensive but concise (about 10-15 sentences).
    """
    response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
    return response.text.strip()

def generate_blog(topic: str, summary: str = "", tone: str = "professional", word_target: int = 700) -> BlogOutput:
    prompt = ChatPromptTemplate.from_template(
        """
        You are an expert content writer. Write clear, engaging, and original long-form blogs.

        Generate a long-form blog post ~{word_target} words about this topic: '{topic}'.
        context: {summary}
        Tone: {tone}.

        Requirements:
        - Provide output in JSON with keys: title, subtitle, body, seo_meta
        - Title: short, engaging (<=12 words)
        - Subtitle: catchy (<=20 words)
        - Body: split into multiple sections with headings
        - Tags: top five key for this blog
        - SEO Meta: concise description (<=160 chars)

        Format:
        {{
            "title": "short, engaging (<=12 words)",
            "subtitle": "catchy (<=20 words)",
            "body": [
                {{"heading": "Section Heading", "content": "Paragraph text"}},
                ...
            ],
            "conclusion": "typically makes up around 10% of the total document's word count, or about 5-7% for theses and dissertations.",
            "tags": ["keyword1", "keyword2"],
            "seo_meta": "meta description (<=160 chars)"
        }}

        Return valid JSON ONLY.
        """
        )
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", api_key=GOOGLE_API_KEY).with_structured_output(BlogOutput)
    try:
        chain = prompt | model
        print(f"📝 Generating blog for topic: {topic}")
        blog = chain.invoke({"topic": topic, "summary": summary, "tone": tone, "word_target": word_target})
        return blog
    except ValidationError as e:
        print("❌ Validation failed:", e)
        raise
    except Exception:
        return BlogOutput(
            title=topic,
            subtitle="",
            body=[{"heading": "Generated Content", "content": "Blog generation failed, please retry."}],
            conclusion="",
            tags=[],
            seo_meta=""
        )
    
# def linkedIn_blog(blog)->str:
#     prompt=f"Make a LinkedIn version for this blog:\n\n{blog}. Strictly respond the content nothing else. Write the blog with proper linkedIn format as linkedin is not support the * as as bolder to use other font style that it looks like it is bolded. and don't need to write here is the linkedin optimised content or any starting statement."
#     response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
#     return response.text.strip()

def linkedIn_blog(blog) -> str:
    """
    Convert blog content to LinkedIn-optimized format using Unicode styling
    """
    prompt = f"""
    Convert this blog content into a LinkedIn-optimized post that will perform well on the platform.
    
    BLOG CONTENT:
    {blog}
    
    Requirements:
    1. Use Unicode characters for formatting (mathematical bold for emphasis, etc.)
    2. Keep it professional yet engaging
    3. Structure with clear sections and bullet points using Unicode symbols
    4. Include relevant hashtags at the end
    5. Optimize for readability and engagement
    6. Length: 800-1200 characters ideal for LinkedIn
    7. Start with a strong hook
    8. Use spacing and emojis strategically (1-2 emojis max)
    9. Include a call-to-action
    
    Formatting Guidelines:
    - Use mathematical bold script (𝗹𝗶𝗸𝗲 𝘁𝗵𝗶𝘀) for key points and headings
    - Use bullet points with • symbol
    - Use proper line breaks for readability
    
    Return ONLY the formatted LinkedIn post content, no explanations.
    """
    
    try:
        response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
        content = response.text.strip()
        
        # Ensure we have proper formatting
        if not any(char in content for char in ['𝗔', '𝘈', '•', '#']):
            # If no formatting detected, add basic structure
            lines = content.split('\n')
            formatted_lines = []
            for i, line in enumerate(lines):
                if i == 0:  # First line - make it bold
                    formatted_lines.append(f"𝗔{line[1:]}𝗔" if line.startswith('A') else f"𝗔{line}𝗔")
                elif line.strip() and not line.startswith('#'):
                    formatted_lines.append(f"• {line}")
                else:
                    formatted_lines.append(line)
            content = '\n'.join(formatted_lines)
        
        return content
        
    except Exception as e:
        print(f"❌ LinkedIn conversion failed: {e}")
        # Fallback basic formatting
        if isinstance(blog, str):
            return f"📝 {blog[:500]}... #Blog #Content"
        else:
            return "🚀 New insights on trending topics! Check out the full blog for more details. #ContentMarketing #Blogging"
        
def export_blog_to_docx(blog, filename):
    doc = Document()

    doc.add_heading(blog.title, 0)
    if blog.subtitle:
        doc.add_paragraph(blog.subtitle, style="Subtitle")

    doc.add_paragraph("")

    for section in blog.body:
        if section.heading:
            doc.add_heading(section.heading, 2)
        if section.content:
            doc.add_paragraph(section.content)

        doc.add_paragraph("")

    if hasattr(blog, "conclusion") and blog.conclusion:
        doc.add_heading("Conclusion", 2)
        doc.add_paragraph(blog.conclusion)
        doc.add_paragraph("")

    if blog.seo_meta:
        p = doc.add_paragraph(blog.seo_meta)
        run = p.runs[0]
        run.italic = True
        doc.add_paragraph("")
        
    if blog.tags:
        doc.add_heading("Tags", 2)
        for tag in blog.tags:
            doc.add_paragraph(tag, style="List Bullet")

    # safe_filename = blog.title.replace(" ", "_").replace("/", "_").replace(":", "_")
    # filepath = blogs_folder / f"{safe_filename}.docx"
    # # filename = safe_filename + ".docx"
    # doc.save(filepath)
    # # doc.save(filename)
    # # print(f"✅ Blog exported to {filename}")
    # print(f"✅ Blog exported to {filepath}")
    safe_filename = re.sub(r'[<>:"/\\|?*]', '_', blog.title)
    safe_filename = safe_filename.replace(" ", "_")
    filepath = blogs_folder / f"{safe_filename}.docx"

    doc.save(filepath)
    print(f"✅ Blog exported to {filepath}")


@app.get("/")
async def root():
    """Root endpoint that returns a hello message"""
    return {"message": "Oh you clever boy!", "status": "success", "service": "Blog Generator API"}

@app.get("/search-topics")
async def search_topics(query: str = Query(..., description="Search query for trending blogs")):
    """Search for topics and return results for user selection"""
    results = google_search(query)
    topics = [{"title": r.get("title"), "url": r.get("link")} for r in results]
    
    if not topics:
        raise HTTPException(status_code=404, detail="No search results found")
    
    # Store search results temporarily
    search_id = str(uuid.uuid4())
    search_store[search_id] = topics
    
    return {
        "search_id": search_id,
        "topics": topics,
        "count": len(topics)
    }

@app.post("/generate-blog-from-topic")
async def generate_blog_from_topic(
    search_id: str,
    selected_topic_index: int = Query(..., ge=0, description="Index of the selected topic"),
    tone: str = Query("professional", description="Tone of the blog"),
    word_target: int = Query(700, description="Target word count")
):
    """Generate blog from user-selected topic"""
    if search_id not in search_store:
        raise HTTPException(status_code=404, detail="Search session not found")
    
    topics = search_store[search_id]
    
    if selected_topic_index >= len(topics):
        raise HTTPException(status_code=400, detail="Invalid topic index")
    
    chosen_topic = topics[selected_topic_index]
    chosen_item = TopicDetails(title=chosen_topic["title"], url=chosen_topic["url"])
    
    summary = summarize_blog(chosen_item.url)
    blog = generate_blog(chosen_item.title, summary, tone, word_target)
    linkedIn_version = linkedIn_blog(blog)
    
    safe_filename = blog.title.replace(" ", "_").replace("/", "_")
    filepath = f"{safe_filename}.docx"
    export_blog_to_docx(blog, safe_filename)
    
    blog_dict = blog.model_dump()
    
    # Clean up stored search results
    if search_id in search_store:
        del search_store[search_id]
    
    return {
        "chosen_topic": chosen_item,
        "summary": summary,
        "docx_file": filepath,
        "blog_content": blog_dict,
        "linkedIn_version": linkedIn_version
    }

@app.post("/generate-blended-blog")
async def generate_blended_blog(request: MultiBlogRequest):
    """Generate a blog by blending multiple selected topics"""
    print("------1")
    # if request.search_id not in search_store:
    #     print("------2")
    #     raise HTTPException(status_code=404, detail="Search session not found")
    
    topics = search_store[request.search_id]
    
    print("------3")
    # Validate all topic indices
    for index in request.topic_indices:
        print("------4")
        if index >= len(topics):
            print("------5")
            raise HTTPException(status_code=400, detail=f"Invalid topic index: {index}")
    
    print("------6")
    # Get selected topics
    selected_topics = [topics[i] for i in request.topic_indices]
    
    print("------7")
    # Summarize each blog
    summaries = []
    print("------8")
    for topic in selected_topics:
        print("------9")
        summary = summarize_blog(topic["url"])
        print("------10")
        summaries.append(summary)
    
    print("------11")
    # Blend summaries
    blended_summary = blend_multiple_blogs(
        summaries, 
        [topic["title"] for topic in selected_topics]
    )
    print("------12")
    
    # Create a combined topic title
    combined_title = f"Comprehensive Analysis: {', '.join([topic['title'][:30] for topic in selected_topics[:3]])}"
    print("------13")
    if len(selected_topics) > 3:
        print("------14")
        combined_title += " and more"
    
    # Generate blog from blended content
    blog = generate_blog(combined_title, blended_summary, request.tone, request.word_target)
    linkedIn_version = linkedIn_blog(blog)
    
    safe_filename = f"blended_blog_{uuid.uuid4().hex[:8]}.docx"
    filepath = safe_filename
    export_blog_to_docx(blog, safe_filename)
    
    # Clean up stored search results
    if request.search_id in search_store:
        del search_store[request.search_id]
    
    return {
        "selected_topics": selected_topics,
        "blended_summary": blended_summary,
        "docx_file": filepath,
        "blog_content": blog.model_dump(),
        "linkedIn_version": linkedIn_version,
        "individual_summaries": summaries
    }

# Keep original endpoint for backward compatibility
@app.get("/generate-blog")
def generate_blog_endpoint(
    query: str = Query(..., description="Search query for trending blogs"), 
    tone: str = Query("professional", description="Tone of the blog"),
    word_target: int = Query(700, description="Target word count")
):
    results = google_search(query)
    topics = [{"title": r.get("title"), "url": r.get("link")} for r in results]
    if not topics:
        return {"error": "No search results found"}

    chosen_item = choose_topic_auto(topics)
    if not chosen_item or not chosen_item.title:
        return {"error": "Gemini did not return a valid topic"}

    summary = summarize_blog(chosen_item.url)
    print("==============summary", summary)
    blog = generate_blog(chosen_item.title, summary, tone, word_target)
    linkedIn_version = linkedIn_blog(blog)
    safe_filename = blog.title.replace(" ", "_").replace("/", "_")
    filepath = f"{safe_filename}.docx"
    export_blog_to_docx(blog, safe_filename)
    blog_dict = blog.model_dump()
    return {
        "topics": topics,
        "chosen_topic": chosen_item,
        "summary": summary,
        "docx_file": filepath,
        "blog_content": blog_dict,
        "linkedIn_version": linkedIn_version
    }

@app.post("/generate-auto-blog")
async def generate_auto_blog(
    search_id: str,
    tone: str = "professional",
    word_target: int = 700
):
    """Automatically choose the best topic and generate blog"""
    if search_id not in search_store:
        raise HTTPException(status_code=404, detail="Search session not found")
    
    topics = search_store[search_id]
    
    # Use the existing auto-topic selection logic
    chosen_item = choose_topic_auto(topics)
    
    # Generate summary and blog
    summary = summarize_blog(chosen_item.url)
    blog = generate_blog(chosen_item.title, summary, tone, word_target)
    linkedIn_version = linkedIn_blog(blog)
    
    safe_filename = blog.title.replace(" ", "_").replace("/", "_")
    filepath = f"{safe_filename}.docx"
    export_blog_to_docx(blog, safe_filename)
    
    blog_dict = blog.model_dump()
    
    # Clean up stored search results
    if search_id in search_store:
        del search_store[search_id]
    
    return {
        "chosen_topic": chosen_item,
        "summary": summary,
        "docx_file": filepath,
        "blog_content": blog_dict,
        "linkedIn_version": linkedIn_version,
        "auto_selected": True
    }

@app.get("/blogs")
def list_blogs():
    blog_files = [f.name for f in blogs_folder.glob("*.docx")]
    return {"blogs": blog_files}

@app.post("/refresh-section")
async def refresh_section(request: dict):
    """
    Refresh or generate content for a new section heading using AI, 
    considering the context of the existing blog.
    Expected JSON input:
    {
        "title": "Blog Title",
        "existing_body": [{"heading": "Intro", "content": "..."}, ...],
        "new_heading": "AI in Healthcare Trends"
    }
    """
    try:
        title = request.get("title", "")
        existing_body = request.get("existing_body", [])
        new_heading = request.get("new_heading", "")

        if not new_heading:
            raise HTTPException(status_code=400, detail="Missing 'new_heading' in request")

        # Build contextual prompt
        context_text = "\n".join(
            [f"{sec['heading']}: {sec.get('content', '')}" for sec in existing_body if sec.get('heading')]
        )
        prompt = f"""
        You are an expert content writer. The blog is titled: "{title}".

        Existing sections:
        {context_text}

        A new section has been added with heading: "{new_heading}".
        Write 2-3 engaging paragraphs for this new section, keeping tone and context consistent.
        Avoid repeating existing information.
        """

        response = genai.GenerativeModel("gemini-2.5-flash").generate_content(prompt)
        generated_text = response.text.strip()

        return {"heading": new_heading, "generated_content": generated_text}

    except Exception as e:
        print(f"❌ Refresh section failed: {e}")
        raise HTTPException(status_code=500, detail=f"AI refresh failed: {e}")

class BlogRequest(BaseModel):
    blog_content: str

@app.post("/generate_instagram_carousel")
def generate_instagram_carousel_endpoint(req: BlogRequest):
    try:
        # Step 1: Generate carousel text
        carousel_text = generate_instagram_carousel(req.blog_content)
        slides = extract_slides(carousel_text)

        # Step 2: Generate images for each slide
        image_files = []
        for i, (title, content) in enumerate(slides.items(), start=1):
            try:
                filename = generate_slide_image(content.strip(), i)
                image_files.append(filename)
            except Exception as e:
                print(f"⚠️ Error generating slide {i}: {e}")
        
        return {
            "carousel_text": carousel_text,
            "images": image_files
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Carousel generation failed: {e}")
    
if __name__ == '__main__':
    uvicorn.run(app, host="0.0.0.0", port=8000)